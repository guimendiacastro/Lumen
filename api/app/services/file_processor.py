# lumen/api/app/services/file_processor.py
"""
Simplified file processor - only handles text extraction
Chunking and embeddings are now handled by RAGService

Enhanced with OCR support for scanned PDFs using Azure Document Intelligence
"""

import io
import logging
from typing import Optional
from dataclasses import dataclass

import PyPDF2
from docx import Document as DocxDocument

from app.services.azure_ocr_service import get_ocr_service, AzureOCRService

log = logging.getLogger("lumen.file_processor")

MAX_DIRECT_CONTEXT_CHARS = 50000


@dataclass
class FileProcessingResult:
    """Result of file processing."""
    use_direct_context: bool
    full_text: str
    total_size: int


class FileProcessor:
    """Simplified file processor - only extracts text"""
    
    @staticmethod
    async def extract_text_from_pdf(content: bytes) -> str:
        """
        Extract text from PDF using PyPDF2 with OCR fallback for scanned documents.

        Strategy:
        1. Try PyPDF2 text extraction first (fast, works for text-based PDFs)
        2. Check if extracted text is sufficient (avg chars per page)
        3. If insufficient (likely scanned), use Azure Document Intelligence OCR

        Args:
            content: PDF file content as bytes

        Returns:
            Extracted text from the PDF
        """
        try:
            # Step 1: Try PyPDF2 extraction first
            pdf_file = io.BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            extracted_text = '\n\n'.join(text_parts)
            num_pages = len(reader.pages)

            # Step 2: Check if OCR is needed
            # Heuristic: If less than 50 chars per page on average, likely scanned
            avg_chars_per_page = len(extracted_text) / max(num_pages, 1)
            needs_ocr = avg_chars_per_page < 50

            if needs_ocr:
                log.info(
                    f"PDF appears to be scanned (avg {avg_chars_per_page:.1f} chars/page). "
                    f"Attempting OCR extraction..."
                )

                # Check if OCR is available
                if not AzureOCRService.is_ocr_available():
                    log.warning(
                        "OCR is needed but Azure Document Intelligence is not configured. "
                        "Returning PyPDF2 text (may be incomplete)."
                    )
                    return extracted_text

                # Step 3: Use OCR for scanned documents
                try:
                    ocr_service = get_ocr_service()
                    ocr_text = await ocr_service.extract_text_with_ocr(content)
                    log.info(
                        f"OCR extraction successful. "
                        f"Extracted {len(ocr_text)} characters from {num_pages} pages."
                    )
                    return ocr_text
                except Exception as ocr_error:
                    log.error(f"OCR extraction failed: {str(ocr_error)}")
                    log.warning("Falling back to PyPDF2 text (may be incomplete)")
                    return extracted_text
            else:
                log.info(
                    f"PDF has sufficient text (avg {avg_chars_per_page:.1f} chars/page). "
                    f"Using PyPDF2 extraction."
                )
                return extracted_text

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    async def extract_text_from_file(content: bytes, mime_type: str) -> str:
        """Extract text from file based on mime type (async to support OCR)"""
        if 'pdf' in mime_type.lower():
            return await FileProcessor.extract_text_from_pdf(content)
        elif 'word' in mime_type.lower() or 'docx' in mime_type.lower():
            return FileProcessor.extract_text_from_docx(content)
        else:
            # Try decoding as text
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                raise ValueError(f"Unsupported file type: {mime_type}")

    @staticmethod
    async def process_file(content: bytes, mime_type: str) -> FileProcessingResult:
        """
        Process file and determine if direct context or RAG should be used.
        Now supports async OCR for scanned PDFs.
        """
        # Extract text (may use OCR for scanned PDFs)
        text = await FileProcessor.extract_text_from_file(content, mime_type)
        text_size = len(text)

        # Decide strategy based on size
        use_direct = text_size <= MAX_DIRECT_CONTEXT_CHARS

        return FileProcessingResult(
            use_direct_context=use_direct,
            full_text=text,
            total_size=text_size
        )